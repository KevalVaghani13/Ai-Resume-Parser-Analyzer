import re
import shutil
import pdfplumber
from docx import Document
import spacy

nlp = spacy.load("en_core_web_sm")

SKILLS_DB = [
    # Languages
    "python", "javascript", "Java", "php", "typescript", "SQL", "C++", "C#", "Ruby", "go", "RUST", "C", "kotlin", "swift", "scala", "dart", "HTML5", "CSS3",

    # Frontend
    "react", "react.js", "react js", "vue", "vue.js", "angular", "html", "css", "bootstrap", "tailwind", "tailwind css", "sass", "next.js", "next js",
    "angular.js", "ember.js", "jquery", "angular js", "ember js", "figma", "canva", "ui/ux", "responsive design", "web design",
    "frontend development", "Responsive Web Design", "Wireframing", "Prototyping", "User Research", "Adobe XD", "Framer", "webflow",

    # Backend
    "node", "node.js", "express", "express.js", "django", "flask", "fastapi", "spring boot", "laravel", "asp.net", ".net",
    "mongodb", "mysql", "postgresql", "sqlite", "redis", "cassandra", "dynamodb", "oracle", "mariadb", "firebase",

    # DevOps & Cloud
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform", "ansible", "cloud computing",
    "git", "github", "gitlab", "bitbucket", "ci/cd",

    # Tools & Platforms
    "postman", "vs code", "visual studio code", "jira", "confluence", "slack", "trello",
    "Adobe Photoshop", "Adobe Illustrator", "CorelDRAW", "Sketch", "InVision",
    "Adobe After Effects", "Adobe Premiere Pro",
    "System Design", "Software Design", "Data Structures", "Algorithms",
    "Object-Oriented Programming", "Software Architecture",
    "Photoshop", "Illustrator", "After Effects", "Premiere Pro",

    # Data & ML
    "machine learning", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
    "data analysis", "nlp", "computer vision",

    # Big Data
    "spark", "hadoop", "hive", "airflow", "kafka", "snowflake", "databricks",

    # APIs & Protocols
    "REST API", "graphql", "websocket", "grpc", "soap", "json", "xml",

    # Other
    "agile", "scrum", "authentication", "authorization", "oauth", "microservices",
    "serverless", "testing", "unit testing", "integration testing",
    "web development", "mobile development", "android", "ios",
    "Computer Engineering", "Computer skills", "IT", "Data entry",
    "Basic computer", "Internet Skills", "Microsoft Office", "MS Office",
    "MS Excel", "MS Word", "MS PowerPoint", "Communication", "presentation",
    "teamwork", "time management", "Problem solving", "Problem-solving",
    "Basic Computer Knowledge", "Hardworking", "Disciplined",
    "JWT", "Communication Skills", "Google Tools", "Leadership",
    "Project management", "Critical thinking", "Adaptability",
    "Collaboration", "Creativity", "Attention to detail",
    "Organizational skills", "Multitasking",
    "Analytical skills", "Time management skills", "Lead Management",
    "Data Management", "Relationship Building Skills",
    "Relationship Building", "JWT Authentication",
    "Cross-Device Compatibility", "Collaboration Skills",
    "Typing Skills", "Customer Convesing Power Skills",
    "team management", "Logo Design", "Brand Identity", "Brand Strategy", "Visual Communication", "Typography", "Color Theory", "UX Research", "Information Architecture", "Design Systems", "User Flows", "Interaction Design", "Usability Testing", "A/B Testing", "Design Thinking", "Empathy", "User-Centered Design", "Mobile-First Design", "Accessibility", "Inclusive Design", "Design Tools",
]


SECTION_ALIASES = {
    "summary": {
        "summary", "profile", "about", "objective", "career objective",
        "professional summary", "summary profile"
    },
    "contact": {
        "contact", "contact me", "contact information", "personal details",
        "reach me", "address", "details"
    },
    "skills": {
        "skills", "technical skills", "core skills", "key skills",
        "technical expertise", "competencies", "skill set", "tech stack",
        "strengths"
    },
    "experience": {
        "experience", "work experience", "employment history", "work history",
        "professional experience", "professional background", "internship",
        "internships", "career history", "work"
    },
    "education": {
        "education", "academic background", "qualification", "qualifications",
        "academics", "academical", "scholarship"
    },
    "projects": {
        "projects", "project experience", "academic projects", "personal projects"
    },
    "certifications": {
        "certifications", "certificates", "licenses", "training"
    },
    "languages": {
        "languages", "language proficiency", "language", "linguistic skills"
    },
    "awards": {
        "awards", "honors", "honours", "achievements", "recognition"
    },
}


def _normalize_text(text):
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = text.replace("\t", " ")
    text = re.sub(r"[ \u00A0]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_line_key(line):
    return re.sub(r"[^a-z0-9]+", " ", line.lower()).strip()


def _normalize_line(line):
    return re.sub(r"\s+", " ", line or "").strip()


def _merge_text_variants(*texts):
    seen = set()
    merged_lines = []

    for text in texts:
        normalized = _normalize_text(text)
        if not normalized:
            continue

        for line in normalized.split("\n"):
            clean_line = _normalize_line(line)
            if not clean_line:
                continue
            line_key = _normalize_line_key(clean_line)
            if line_key in seen:
                continue
            seen.add(line_key)
            merged_lines.append(clean_line)

    return "\n".join(merged_lines)


def _section_name_for_line(line):
    line_key = _normalize_line_key(line)
    if not line_key:
        return None

    for canonical, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            alias_key = _normalize_line_key(alias)
            if line_key == alias_key:
                return canonical
            if line_key.startswith(alias_key + " "):
                return canonical
            if line_key.endswith(" " + alias_key):
                return canonical
    return None


def extract_section_map(text):
    normalized = _normalize_text(text)
    lines = [line for line in (_normalize_line(line) for line in normalized.split("\n")) if line]
    sections = {name: [] for name in SECTION_ALIASES}
    sections["_preamble"] = []

    current = "_preamble"

    for line in lines:
        inline_match = re.match(r"^([A-Za-z][A-Za-z /&+\-]{1,40}?)(?:\s*[:\-]\s*)(.+)$", line)
        if inline_match:
            header = _normalize_line(inline_match.group(1))
            remainder = _normalize_line(inline_match.group(2))
            section_name = _section_name_for_line(header)
            if section_name:
                current = section_name
                if remainder:
                    sections[current].append(remainder)
                continue

        section_name = _section_name_for_line(line)
        if section_name and len(line.split()) <= 6:
            current = section_name
            continue

        sections.setdefault(current, []).append(line)

    sections["__all__"] = lines
    return sections


def _section_text(section_map, *section_names):
    lines = []
    for name in section_names:
        lines.extend(section_map.get(name, []))
    return "\n".join(lines).strip()


def _split_words_into_columns(words, page_width, depth=0):
    if depth >= 2 or len(words) < 12:
        return [words]

    x_positions = sorted(
        {
            round(float(word.get("x0", 0)), 1)
            for word in words
            if _normalize_line(word.get("text", ""))
        }
    )
    if len(x_positions) < 6:
        return [words]

    gap_candidates = [
        (x_positions[index + 1] - x_positions[index], index)
        for index in range(len(x_positions) - 1)
    ]
    if not gap_candidates:
        return [words]

    largest_gap, gap_index = max(gap_candidates, key=lambda item: item[0])
    threshold = max(80.0, (page_width or 0) * 0.18)
    if largest_gap < threshold:
        return [words]

    split_x = (x_positions[gap_index] + x_positions[gap_index + 1]) / 2.0
    left_words = [word for word in words if float(word.get("x0", 0)) <= split_x]
    right_words = [word for word in words if float(word.get("x0", 0)) > split_x]

    if len(left_words) < 4 or len(right_words) < 4:
        return [words]

    left_columns = _split_words_into_columns(left_words, page_width, depth + 1)
    right_columns = _split_words_into_columns(right_words, page_width, depth + 1)
    return left_columns + right_columns


def _reconstruct_text_from_words(words):
    if not words:
        return ""

    sorted_words = sorted(words, key=lambda item: (float(item.get("top", 0)), float(item.get("x0", 0))))
    lines = []
    current_line = []
    current_top = None
    tolerance = 3.0

    for word in sorted_words:
        text = _normalize_line(word.get("text", ""))
        if not text:
            continue

        top = float(word.get("top", 0))
        if current_top is None or abs(top - current_top) <= tolerance:
            current_line.append(word)
            current_top = top if current_top is None else (current_top + top) / 2.0
        else:
            line = " ".join(
                _normalize_line(item.get("text", ""))
                for item in sorted(current_line, key=lambda item: float(item.get("x0", 0)))
            )
            if line:
                lines.append(line)
            current_line = [word]
            current_top = top

    if current_line:
        line = " ".join(
            _normalize_line(item.get("text", ""))
            for item in sorted(current_line, key=lambda item: float(item.get("x0", 0)))
        )
        if line:
            lines.append(line)

    return "\n".join(lines)


def _extract_page_words_text(page):
    try:
        words = page.extract_words(keep_blank_chars=False, use_text_flow=True)
    except TypeError:
        words = page.extract_words()
    except Exception:
        return ""

    if not words:
        return ""

    page_width = float(getattr(page, "width", 0) or 0)
    column_texts = []
    for column_words in _split_words_into_columns(words, page_width):
        column_text = _reconstruct_text_from_words(column_words)
        if column_text.strip():
            column_texts.append(column_text)

    return "\n\n".join(column_texts)


def _experience_context_text(text):
    section_map = extract_section_map(text)
    experience_text = _section_text(section_map, "experience")
    if experience_text:
        return experience_text

    return _normalize_text(text)


def _is_sparse_text(text):
    normalized = _normalize_text(text)
    if not normalized:
        return True

    alnum_count = len(re.findall(r"[A-Za-z0-9]", normalized))
    return alnum_count < 80 or len(normalized) < 120


def _ocr_pdf_page(pdf_path, page_index, zoom=2.5):
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    # Ensure tesseract binary is available on PATH (pytesseract is just a wrapper)
    if shutil.which("tesseract") is None:
        return ""

    ocr_configs = ["--psm 11", "--psm 12", "--psm 6"]

    def _ocr_image(image):
        outputs = []
        for config in ocr_configs:
            try:
                text = pytesseract.image_to_string(image, config=config)
            except Exception:
                text = ""
            if text and text.strip():
                outputs.append(text)
        return _merge_text_variants(*outputs)

    # Prefer PyMuPDF (fitz) for high-quality rendering when available
    try:
        import fitz
        doc = fitz.open(pdf_path)
        try:
            page = doc.load_page(page_index)
            matrix = fitz.Matrix(max(zoom, 3.5), max(zoom, 3.5))
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)

            image_mode = "RGBA" if pixmap.n >= 4 else "RGB"
            image = Image.frombytes(image_mode, [pixmap.width, pixmap.height], pixmap.samples)
            if image_mode == "RGBA":
                image = image.convert("RGB")

            ocr_text = _ocr_image(image)
            return ocr_text or ""
        except Exception:
            return ""
        finally:
            try:
                doc.close()
            except Exception:
                pass
    except Exception:
        # If fitz is not available or fails, fall back to pdfplumber rendering
        try:
            with pdfplumber.open(pdf_path) as pdf:
                if page_index < 0 or page_index >= len(pdf.pages):
                    return ""
                page = pdf.pages[page_index]
                try:
                    pil_img = page.to_image(resolution=150).original
                except Exception:
                    # Last resort: try converting bbox snapshot
                    try:
                        pil_img = page.to_image().original
                    except Exception:
                        return ""

                ocr_text = _ocr_image(pil_img)
                return ocr_text or ""
        except Exception:
            return ""


def extract_text_from_pdf(path):
    page_texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            variants = [
                page.extract_text() or "",
                page.extract_text(layout=True) or "",
                _extract_page_words_text(page),
            ]
            page_text = _merge_text_variants(*variants)

            needs_ocr = _is_sparse_text(page_text) or not re.search(r'\b\d{10}\b', page_text)
            if needs_ocr:
                ocr_text = _ocr_pdf_page(path, page.page_number - 1)
                if ocr_text:
                    page_text = _merge_text_variants(page_text, ocr_text)

            if page_text:
                page_texts.append(page_text)

    return _merge_text_variants(*page_texts)


def extract_text_from_docx(path):
    doc = Document(path)
    chunks = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

    return _merge_text_variants("\n".join(chunks))


def extract_skills(text):
    """
    Extract skills from text using custom boundary detection.
    Standard \b fails for skills like C++, C#, .NET because +, #, . are not word characters.
    """
    normalized = _normalize_text(text)
    lines = [_normalize_line(line) for line in normalized.split('\n') if _normalize_line(line)]
    section_map = extract_section_map(text)
    skills_text = _section_text(section_map, "skills", "languages")
    ambiguous_skills = {"c", "go", "r", "it"}

    canonical_skill_map = {
        "photoshop": "Adobe Photoshop",
        "adobe photoshop": "Adobe Photoshop",
        "illustrator": "Adobe Illustrator",
        "adobe illustrator": "Adobe Illustrator",
        "after effects": "Adobe After Effects",
        "adobe after effects": "Adobe After Effects",
        "react.js": "react",
        "react js": "react",
        "node.js": "node.js",
        "node": "node.js",
    }

    def _scan_skills(source_text, strict_ambiguous):
        text_lower = _normalize_text(source_text).lower()
        if not text_lower:
            return []

        hard_noise = {
            "table", "invoice", "amount", "menu", "services", "website link",
            "application link", "post link", "contact us", "instagram", "logo"
        }
        if any(noise in text_lower for noise in hard_noise):
            text_lower = re.sub(r"\b(?:table|invoice|amount|menu|services|instagram|contact us)\b", " ", text_lower)

        found = []
        for skill in SKILLS_DB:
            skill_lower = skill.lower()

            if skill_lower in ambiguous_skills and not strict_ambiguous:
                continue

            pattern = r'(?:^|(?<=[^a-zA-Z0-9]))' + re.escape(skill_lower) + r'(?=[^a-zA-Z0-9]|$)'
            if skill_lower in ambiguous_skills and strict_ambiguous:
                strict_pattern = r'(?:^|[\s,|/])' + re.escape(skill_lower) + r'(?=$|[\s,|/])'
                if not re.search(strict_pattern, text_lower):
                    continue

            if re.search(pattern, text_lower):
                found.append(skill)

        return found

    def _skills_window_text_from_headings(all_lines):
        section_stoppers = {
            "summary", "about", "objective", "experience", "education",
            "projects", "certifications", "awards", "portfolio", "design showcase"
        }
        captured = []
        for index, line in enumerate(all_lines):
            if _normalize_line_key(line) != "skills":
                continue

            for look_ahead in range(index + 1, min(index + 26, len(all_lines))):
                candidate = all_lines[look_ahead]
                candidate_key = _normalize_line_key(candidate)

                if candidate_key in section_stoppers:
                    break
                if candidate_key == "contact":
                    continue
                if re.search(r'@|https?://|www\.', candidate, re.IGNORECASE):
                    continue
                if re.search(r'\+?\d[\d\s().-]{8,}\d', candidate):
                    continue
                if len(candidate.split()) > 6:
                    continue

                captured.append(candidate)

        return "\n".join(captured)

    # 1) best effort around literal SKILLS headings (robust for interleaved OCR columns)
    heading_window_text = _skills_window_text_from_headings(lines)
    heading_found = _scan_skills(heading_window_text, strict_ambiguous=True)

    # 2) structured section map fallback
    structured_found = _scan_skills(skills_text, strict_ambiguous=True)

    # 3) full-text fallback only when section-like extraction fails
    full_found = _scan_skills(text, strict_ambiguous=False)

    if heading_found and len(heading_found) >= 2:
        skills_found = heading_found
    elif structured_found and len(structured_found) >= 2:
        skills_found = structured_found
    else:
        skills_found = full_found

    # If skills came from section context, don't inject role-title tags like ui/ux unless present there.
    section_context_lower = _normalize_text(heading_window_text + "\n" + skills_text).lower()
    if section_context_lower and "ui/ux" not in section_context_lower:
        skills_found = [skill for skill in skills_found if skill.lower() != "ui/ux"]

    deduped = []
    seen = set()
    for skill in skills_found:
        canonical = canonical_skill_map.get(skill.lower(), skill)
        key = canonical.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(canonical)

    return deduped


def extract_keywords_from_job(jd_text):
    """Extract focused, high-signal keywords from job description."""
    jd_clean = re.sub(r"[\n\r\t]+", " ", jd_text.lower())
    jd_clean = re.sub(r"\s+", " ", jd_clean).strip()

    stop_words = {
        "with", "from", "that", "this", "have", "your", "will", "role", "and",
        "for", "the", "about", "you", "experience", "years", "or", "in", "nice",
        "our", "team", "work", "ability", "strong", "good", "well", "etc", "due",
        "required", "requirements", "developer", "engineer", "senior", "junior",
        "position", "job", "like", "such", "more", "most"
    }

    # Prioritize requirement-heavy lines to improve precision.
    focused_lines = []
    for line in jd_text.splitlines():
        line_l = line.lower().strip()
        if not line_l:
            continue
        if any(key in line_l for key in ["require", "must", "should", "need", "skill", "experience", "responsib", "expert"]):
            focused_lines.append(line_l)

    focused_text = " ".join(focused_lines) if focused_lines else jd_clean

    phrases = re.findall(r"\b[a-z]{3,}(?:\s+[a-z]{3,}){1,2}\b", focused_text)
    single_words = re.findall(r"\b[a-z]{4,}\b", focused_text)

    # Keep known skills from the full JD even if they were outside focused lines.
    explicit_skills = [skill for skill in SKILLS_DB if re.search(r"\b" + re.escape(skill) + r"\b", jd_clean)]

    keywords = set(phrases + single_words + explicit_skills)
    cleaned = []
    for kw in keywords:
        if kw in stop_words:
            continue
        if kw.count(" ") > 2:
            continue
        if any(c in kw for c in ".,;:!?@#$%^&*()"):
            continue
        cleaned.append(kw.strip())

    # Keep the strongest 60 to avoid noisy denominator in scoring.
    cleaned.sort(key=lambda x: (x.count(" "), len(x)), reverse=True)
    return cleaned[:60]


def match_jd_keywords_in_resume(resume_text, jd_keywords):
    """Find which JD keywords appear in resume with efficient token/phrase matching."""
    resume_lower = re.sub(r"\s+", " ", resume_text.lower())
    token_set = set(re.findall(r"\b[a-z0-9\+#\.]+\b", resume_lower))
    matched_keywords = set()

    for kw in jd_keywords:
        kw_norm = kw.strip().lower()
        if not kw_norm:
            continue
        if " " in kw_norm:
            if re.search(r"\b" + re.escape(kw_norm) + r"\b", resume_lower):
                matched_keywords.add(kw_norm)
        else:
            if kw_norm in token_set:
                matched_keywords.add(kw_norm)

    return list(matched_keywords)


def extract_experience_duration(text):
    """Extract the most likely work-experience duration/period from the resume.

    Works with diverse layouts by preferring explicit experience sections when
    available, then falling back to context-aware line scanning.
    """
    section_map = extract_section_map(text)
    experience_text = _section_text(
        section_map,
        "experience",
    ) or _normalize_text(text)

    education_keywords = {
        "education", "school", "college", "university", "ssc", "hsc", "10th",
        "12th", "bachelor", "master", "degree", "diploma", "cgpa", "gpa",
        "bca", "bcom", "bsc", "mca", "mcom", "msc", "be", "b.e.", "btech",
        "mtech", "phd"
    }
    job_keywords = {
        "experience", "work", "intern", "internship", "developer", "engineer",
        "analyst", "associate", "lead", "manager", "freelance", "project",
        "worked", "role", "company", "at"
    }

    month_name = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*"
    duration_patterns = [
        r"\b\d+\+?\s*(?:years?|yrs?|year|yr)\b",
        r"\b\d+\+?\s*(?:months?|mos?|month|mo)\b",
        r"\b\d+\+?\s*(?:weeks?|wks?|week|wk)\b",
        rf"\b{month_name}\.?\s*\d{{4}}\s*[-–]\s*(?:{month_name}|present|current|now)\.?\s*(?:\d{{4}})?\b",
        r"\b(?:19|20)\d{2}\s*[-–]\s*(?:present|current|now|(?:19|20)\d{2})\b",
        r"\b(?:19|20)\d{2}\s*[-–]\s*(?:19|20)\d{2}\b",
    ]

    best_match = None
    best_score = -1

    for raw_line in experience_text.split("\n"):
        line = _normalize_line(raw_line)
        if not line:
            continue

        lower = line.lower()
        if any(keyword in lower for keyword in education_keywords):
            continue

        line_score = 0
        if any(keyword in lower for keyword in job_keywords):
            line_score += 2
        if any(char.isdigit() for char in line):
            line_score += 1

        for pattern in duration_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                candidate = _normalize_line(match.group(0))
                score = line_score
                if re.search(r"\b(?:years?|yrs?|months?|mos?|weeks?|wks?)\b", candidate, re.IGNORECASE):
                    score += 3
                if re.search(r"\b(?:present|current|now)\b", candidate, re.IGNORECASE):
                    score += 2
                if len(candidate) > 2:
                    score += 1

                if score > best_score:
                    best_score = score
                    best_match = candidate

    return best_match


def extract_name(text):
    """Extract candidate name across diverse resume layouts.

    Tries explicit labels first, then spaced-letter headers, then scores the
    top header lines before falling back to spaCy PERSON entities.
    """
    normalized = _normalize_text(text)
    lines = [line.strip() for line in normalized.split('\n') if line.strip()]
    if not lines:
        return None

    label_patterns = [
        re.compile(r'^\s*name\s*[:\-]\s*(.+)$', re.IGNORECASE),
        re.compile(r'^\s*candidate\s+name\s*[:\-]\s*(.+)$', re.IGNORECASE),
        re.compile(r'^\s*full\s+name\s*[:\-]\s*(.+)$', re.IGNORECASE),
        re.compile(r'my\s+name\s+is\s+([A-Za-z\s]+?)(?:\s+and|\s+,|$)', re.IGNORECASE),  # "My Name is XXX"
    ]

    section_index = _first_section_index(lines)
    header_limit = section_index if section_index is not None else min(len(lines), 12)
    header_lines = lines[:max(1, header_limit)]
    name_search_lines = lines[:min(20, len(lines))]

    # 1. Explicit name labels like "Name : Amisha Subhashbhai Gavit" or "My Name is Priyanshi".
    for line in header_lines:
        for pattern in label_patterns:
            match = pattern.search(line)
            if match:
                candidate = _normalize_line(match.group(1))
                candidate = _collapse_spaced_letters(candidate)
                candidate = _expand_single_word_name(candidate, name_search_lines)
                if _name_score(candidate) >= 3:
                    return _prettify_name(candidate.title() if candidate.isupper() else candidate)

    # 2. Spaced-letter name blocks like "M I R A L" + "K A C H A R I Y A".
    for index in range(min(len(header_lines) - 1, 14)):
        first = header_lines[index]
        second = header_lines[index + 1]
        if _looks_like_spaced_name_segment(first) and _looks_like_spaced_name_segment(second):
            candidate = f"{_collapse_spaced_letters(first)} {_collapse_spaced_letters(second)}"
            if _name_score(candidate) >= 4:
                return _prettify_name(candidate.title() if candidate.isupper() else candidate)

    # 2b. OCR-specific: Extract name from patterns like "Vidhi sheladiya Contact" or "Priyanshi Vaghani Surat"
    #     where name comes before contact/location keywords.
    #     Search broader (first 15 lines) since OCR layouts vary
    search_lines = lines[:min(15, len(lines))]
    for line in search_lines:
        # Try to extract name before "Contact", "Phone", "Email", "Surat", "India", etc.
        # More flexible pattern to handle OCR variations like lowercase surnames
        ocr_pattern = re.compile(r'^([A-Z][a-z]+(?:\s+[a-z]+)*)\s+(?:Contact|Phone|Mobile|Email|Surat|India|Bangalore|Mumbai|Delhi|Location)', re.IGNORECASE)
        match = ocr_pattern.match(line)
        if match:
            candidate = _normalize_line(match.group(1))
            candidate = _expand_single_word_name(candidate, name_search_lines)
            if _name_score(candidate) >= 2:  # Lower threshold for OCR
                return _prettify_name(candidate.title() if candidate.isupper() else candidate)

    # 2c. OCR layouts sometimes place the Contact heading before the name.
    #     In that case, the line immediately after Contact is often the candidate name.
    contact_labels = {'contact', 'contact me', 'contact information', 'personal details', 'reach me'}
    for index, line in enumerate(search_lines[:-1]):
        if _normalize_line_key(line) in contact_labels:
            for offset in (1, 2):
                next_index = index + offset
                if next_index >= len(search_lines):
                    continue
                candidate = _collapse_spaced_letters(search_lines[next_index])
                candidate = _expand_single_word_name(candidate, name_search_lines)
                if _name_score(candidate) >= 2 and '@' not in candidate and not re.search(r'\d', candidate):
                    return _prettify_name(candidate.title() if candidate.isupper() else candidate)

    # 3. Score the top header lines and keep the most name-like one.
    best_candidate = None
    best_score = -1
    for line in header_lines:
        if _section_name_for_line(line):
            continue
        if any(token in line.lower() for token in ['@', 'http', '.com', '.in', '.net', '.org']):
            continue
        if re.search(r'\d', line):
            continue

        candidate = _collapse_spaced_letters(line)
        score = _name_score(candidate)
        if score > best_score:
            best_candidate = candidate
            best_score = score

    if best_candidate and best_score >= 3:
        best_candidate = _expand_single_word_name(best_candidate, name_search_lines)
        return _prettify_name(best_candidate.title() if best_candidate.isupper() else best_candidate)

    # 4. Fallback to PERSON entities from the header region.
    header_text = "\n".join(header_lines)
    doc = nlp(header_text[:800])
    for ent in doc.ents:
        if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
            return _prettify_name(ent.text)

    # 5. Final fallback: allow a short, clean header line.
    for line in header_lines[:8]:
        candidate = _collapse_spaced_letters(line)
        candidate = _expand_single_word_name(candidate, name_search_lines)
        if _name_score(candidate) >= 2:
            return _prettify_name(candidate.title() if candidate.isupper() else candidate)

    return None


def extract_location(text):
    """
    Extract location (area, city, state, country) from the resume header.
    Prioritizes lines with multiple commas representing a full address.
    """
    normalized = _normalize_text(text)
    lines = [line.strip() for line in normalized.split('\n') if line.strip()]
    if not lines:
        return None

    # Prefer compact address-like lines near the top of the document.
    address_keywords = {
        'address', 'location', 'city', 'state', 'country', 'india', 'pin', 'pincode',
        'surat', 'ahmedabad', 'vadodara', 'rajkot', 'mumbai', 'delhi', 'pune', 'bangalore',
        'hyderabad', 'chennai', 'noida', 'gandhinagar', 'gujarat', 'maharashtra', 'karnataka'
    }
    noise_keywords = {
        'experience', 'education', 'project', 'projects', 'skills', 'summary', 'objective',
        'worked', 'working', 'developed', 'created', 'built', 'implemented', 'designed',
        'intern', 'internship', 'profile', 'contact', 'email', 'phone', 'mobile'
    }

    def _location_score(line: str) -> tuple[int, int, int]:
        lower = line.lower()
        if '@' in lower:
            return (-100, 0, 0)
        if re.search(r'https?://|www\.', lower):
            return (-100, 0, 0)
        # Never treat phone-like lines as location.
        if re.search(r'\+?\d[\d\s().-]{8,}\d', line):
            return (-100, 0, 0)

        score = 0
        if any(keyword in lower for keyword in address_keywords):
            score += 4
        if re.search(r'\b(?:surat|india|gujarat|mumbai|delhi|pune|bangalore|hyderabad|chennai|noida)\b', lower):
            score += 4
        if ',' in line:
            score += min(3, line.count(','))
        if re.search(r'\d{6}', line):
            score += 2
        if len(line) <= 60:
            score += 1
        if any(keyword in lower for keyword in noise_keywords):
            score -= 6
        if not any(keyword in lower for keyword in address_keywords) and ',' not in line and not re.search(r'\d{6}', line):
            score -= 4
        if len(line.split()) > 8:
            score -= 2
        return (score, line.count(','), -len(line))

    search_limit = min(30, len(lines))
    candidate_lines = []

    for i in range(search_limit):
        line = lines[i]
        parts = [part.strip() for part in re.split(r'[|•·*]', line) if part.strip()]
        if parts:
            candidate_lines.extend(parts)
        else:
            candidate_lines.append(line)

    best_candidate = None
    best_score = (-999, 0, 0)

    for line in candidate_lines:
        score = _location_score(line)
        if score > best_score:
            best_score = score
            best_candidate = line

    if best_candidate and best_score[0] >= 4:
        return _normalize_line(best_candidate).rstrip(' ,.-')

    # NER fallback for layouts that do not expose commas cleanly.
    header_doc = nlp("\n".join(lines[:12]))
    gpe_parts = [ent.text for ent in header_doc.ents if ent.label_ in {"GPE", "LOC", "FAC"} and len(ent.text) <= 60]
    if gpe_parts:
        return ", ".join(dict.fromkeys(gpe_parts[:3])).rstrip(' ,.-')

    return None


# Action verbs that indicate strong experience
ACTION_VERBS = {
    'built', 'designed', 'developed', 'created', 'implemented', 'engineered',
    'led', 'managed', 'directed', 'supervised', 'oversaw',
    'optimized', 'improved', 'enhanced', 'streamlined', 'accelerated',
    'delivered', 'deployed', 'released', 'launched', 'shipped',
    'increased', 'grew', 'scaled', 'expanded', 'distributed',
    'reduced', 'decreased', 'minimized', 'eliminated', 'cut',
    'automated', 'integrated', 'migrated', 'refactored', 'architected'
}


def extract_action_verbs(text):
    """Find action verbs that indicate strong experience."""
    text_lower = text.lower()
    found_verbs = []
    for verb in ACTION_VERBS:
        if re.search(r'\b' + verb + r'\b', text_lower):
            found_verbs.append(verb)
    return list(set(found_verbs))


def extract_metrics_and_impact(text):
    """Extract quantifiable metrics and impact statements."""
    # Look for patterns like: "30% improvement", "2x faster", "$5M", "100K users"
    metrics = re.findall(r'\b(\d+[%x]|\$\d+[KMB]?|\d+\s*(?:K|M|B|million|thousand))\b', text, re.IGNORECASE)
    
    # Look for action phrases with numbers
    impact_patterns = [
        r'(?:improved|increased|reduced|saved|generated|earned|achieved)\s+(?:by\s+)?(\d+[%])',
        r'(\d+)\s*(?:x|times|fold)\s*(?:faster|more|better)',
    ]
    
    impact_statements = []
    for pattern in impact_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        impact_statements.extend(matches)
    
    return {
        'metrics_count': len(metrics),
        'has_impact': len(impact_statements) > 0,
        'metric_samples': metrics[:3]
    }


def extract_experience_bullets(text):
    """Extract experience bullets/achievements from text."""
    section_map = extract_section_map(text)
    experience_text = _section_text(section_map, 'experience') or _normalize_text(text)
    lines = experience_text.split('\n')
    bullets = []
    signal_verbs = ACTION_VERBS.union({
        'completed', 'learned', 'worked', 'created', 'built', 'supported',
        'contributed', 'collaborated', 'assisted', 'handled', 'maintained'
    })
    
    # Find lines that start with bullet points or dashes
    for line in lines:
        line_stripped = line.strip()
        if line_stripped.startswith(('•', '-', '*', '○', '–', '—')) or (len(line_stripped) > 20 and any(verb in line_stripped.lower() for verb in signal_verbs)):
            # Clean up the bullet point
            clean_line = re.sub(r'^[•\-*○–—]\s*', '', line_stripped)
            if len(clean_line) > 10:  # minimum useful length
                bullets.append(clean_line)
    
    return bullets[:20]  # Return top 20 bullets


def extract_education(text):
    """Extract education information."""
    education_keywords = {
        'bachelor', 'master', 'phd', 'diploma', 'degree', 'b.s.', 'b.a.', 'm.s.',
        'm.a.', 'b.tech', 'btech', 'bca', 'b.c.a', 'bcom', 'b.com', 'bsc', 'b.sc',
        'mca', 'm.c.a', 'mcom', 'm.com', 'msc', 'm.sc', 'be', 'b.e.',
        'ssc', 'hsc', '10th', '12th', 'cgpa', 'gpa', 'graduation', 'post-graduation',
        'application', 'science', 'gseb', 'cbse', 'icse'
    }
    institute_keywords = {
        'school', 'college', 'university', 'board', 'institute', 'academy',
        'vidhyalay', 'sankul', 'vidhyalaya', 'vidyalaya', 'high school'
    }
    noise_tokens = {
        'figma', 'application link', 'website link', 'post link', 'logo design',
        'tool skills', 'contact', 'project', 'experience', 'services', 'menu',
        'portfolio', 'instagram', 'ui/ux', 'link'
    }

    section_map = extract_section_map(text)
    education_text = _section_text(section_map, 'education')
    # If the section-based extraction is too short or empty, fall back to scanning the entire text.
    source_text = education_text if (education_text and len(education_text.strip()) > 30) else _normalize_text(text)
    lines = [_normalize_line(line) for line in source_text.split('\n') if _normalize_line(line)]

    def _sanitize_edu_line(raw_line):
        line = raw_line
        line = re.sub(r'^[^A-Za-z0-9]+', '', line)
        line = re.sub(r'\s+', ' ', line).strip()

        # Remove trailing OCR/link clutter while preserving core education text.
        split_pattern = r'\b(?:figma|application link|website link|post link|tool skills|logo design|portfolio|contact)\b'
        line = re.split(split_pattern, line, flags=re.IGNORECASE)[0].strip()
        return line.rstrip(' ,;:-')

    education = []
    for raw_line in lines:
        line = _sanitize_edu_line(raw_line)
        if not line:
            continue

        lower = line.lower()
        if any(token in lower for token in noise_tokens):
            continue

        has_degree = any(keyword in lower for keyword in education_keywords)
        has_institute = any(keyword in lower for keyword in institute_keywords)
        years = re.findall(r'\b(?:19|20)\d{2}\b', line)
        has_year = bool(years)

        if not (has_degree or has_institute or (education_text and has_year)):
            continue

        # Avoid saving isolated single years like "2025" as education entries.
        if has_year and not (has_degree or has_institute):
            has_year_range = re.search(r'\b(?:19|20)\d{2}\s*[-–]\s*(?:19|20)\d{2}\b', line) is not None
            if not has_year_range:
                continue
            if education:
                education[-1] = f"{education[-1]} {line}".strip()[:140]
            continue

        # Merge institute line into preceding degree line when they belong together.
        if education:
            prev_lower = education[-1].lower()
            prev_has_degree = any(keyword in prev_lower for keyword in education_keywords)
            prev_has_institute = any(keyword in prev_lower for keyword in institute_keywords)

            if has_institute and prev_has_degree and not prev_has_institute:
                education[-1] = f"{education[-1]} {line}".strip()[:140]
                continue
            if has_degree and prev_has_institute and not prev_has_degree:
                education[-1] = f"{education[-1]} {line}".strip()[:140]
                continue

        education.append(line[:140])

    # Deduplicate while preserving order.
    deduped = []
    seen = set()
    for item in education:
        key = _normalize_line_key(item)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    # Drop short entries that are already contained inside longer education lines.
    pruned = []
    normalized_items = [_normalize_line(item) for item in deduped]
    for index, item in enumerate(normalized_items):
        item_lower = item.lower()
        is_subset = False
        for j, other in enumerate(normalized_items):
            if index == j:
                continue
            other_lower = other.lower()
            if len(item_lower) < len(other_lower) and item_lower in other_lower:
                is_subset = True
                break
        if not is_subset:
            pruned.append(deduped[index])

    return pruned[:6]


def check_formatting_quality(text):
    """Score resume formatting quality."""
    score = 100
    penalties = 0
    
    # Check for standard sections
    standard_sections = ['experience', 'education', 'skills', 'projects', 'summary', 'objective']
    found_sections = sum(1 for section in standard_sections if section in text.lower())
    
    if found_sections < 2:
        penalties += 10
    
    # Penalize excessive special characters
    special_char_count = len(re.findall(r'[™©®]', text))
    if special_char_count > 5:
        penalties += 5
    
    # Check line length consistency (very long lines indicate formatting issues)
    lines = text.split('\n')
    long_lines = sum(1 for line in lines if len(line) > 120)
    if long_lines > len(lines) * 0.3:  # more than 30% very long lines
        penalties += 5
    
    # Reward for clear structure (multiple small sections)
    if found_sections >= 3:
        score += 10
    
    return max(0, min(100, score - penalties))


def extract_email(text):
    """Extract email address and clean it of stray characters."""
    # Robust email regex
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    if match:
        email = match.group(0).strip().lower()
        # Ensure it doesn't end with a dot (common if it's at the end of a sentence)
        return email.rstrip('.')
    return None


NAME_STOPWORDS = {
    'about', 'address', 'contact', 'education', 'experience', 'languages',
    'portfolio', 'projects', 'skills', 'summary', 'objective', 'profile',
    'work', 'internship', 'certifications', 'designation', 'gender', 'hobbies',
    'declaration', 'resume', 'curriculum', 'vitae', 'me', 'my', 'and', 'of',
    'the', 'to', 'for', 'graphic', 'designer', 'developer', 'engineer',
    'student', 'analyst', 'consultant', 'intern', 'professional', 'adobe',
    'photoshop', 'illustrator', 'figma', 'coreldraw', 'effects', 'after', 'ui',
    'ux', 'creative', 'detail', 'oriented', 'communication', 'skills',
    'university', 'college', 'institute', 'school', 'board', 'academy',
    'vidhyalay', 'sankul', 'hsc', 'ssc', 'gseb', 'cbse', 'icse', 'degree',
    'diploma', 'bachelor', 'master', 'phd', 'vnsgu', 'gujarat', 'surat', 'india'
}


def _collapse_spaced_letters(line):
    tokens = [token for token in re.split(r'\s+', line.strip()) if token]
    if len(tokens) >= 2 and all(len(token) == 1 and token.isalpha() for token in tokens):
        return ''.join(tokens)
    return line.strip()


def _looks_like_spaced_name_segment(line):
    tokens = [token for token in re.split(r'\s+', line.strip()) if token]
    return len(tokens) >= 2 and len(tokens) <= 12 and all(len(token) == 1 and token.isalpha() for token in tokens)


def _first_section_index(lines):
    for index, line in enumerate(lines):
        if _section_name_for_line(line):
            return index
    return None


def _prettify_name(name):
    tokens = []
    for token in _normalize_line(name).split():
        if token.islower():
            tokens.append(token.capitalize())
        else:
            tokens.append(token)
    return " ".join(tokens).strip()


def _expand_single_word_name(candidate, search_lines):
    words = [word for word in re.split(r"\s+", _normalize_line(candidate)) if word]
    if len(words) != 1:
        return candidate

    base = words[0].lower()
    best = candidate
    best_score = _name_score(candidate)

    for line in search_lines:
        cleaned = re.sub(r"[^A-Za-z\s]", " ", _normalize_line(line))
        line_words = [word for word in re.split(r"\s+", cleaned) if word]
        if len(line_words) < 2 or len(line_words) > 4:
            continue
        if base not in {word.lower() for word in line_words}:
            continue

        full_candidate = " ".join(line_words)
        score = _name_score(full_candidate)
        if score > best_score:
            best = full_candidate
            best_score = score

    return best


def _name_score(line):
    if not line:
        return -1

    line = _normalize_line(line)
    if not line:
        return -1

    if '@' in line or 'http' in line.lower() or re.search(r'\d{5,}', line):
        return -1
    if ',' in line or '|' in line:
        return -1

    cleaned = re.sub(r'[^A-Za-z\s\.]', ' ', line)
    words = [word for word in re.split(r'\s+', cleaned) if word]
    if not words:
        return -1

    lower_words = {word.lower().strip('.').strip() for word in words}
    if lower_words.intersection(NAME_STOPWORDS):
        return -1

    has_proper_case = line.isupper() or any(word[0].isupper() for word in words if word)
    if not has_proper_case:
        return -1

    # Very short one-word names are possible, but only in a header region and
    # only when they look like an actual proper noun.
    all_alpha = all(word.replace('.', '').isalpha() for word in words)
    if not all_alpha:
        return -1

    score = 0
    if 2 <= len(words) <= 4:
        score += 4
    elif len(words) == 1:
        score += 1
    else:
        score -= 2

    if sum(1 for word in words if word[0].isupper()) >= 1:
        score += 2

    if line.isupper():
        score += 1

    if any(token.lower() in {'mr', 'mrs', 'ms', 'miss', 'dr'} for token in lower_words):
        score += 1

    return score


def _normalize_phone_candidate(raw, has_label=False):
    if not raw:
        return None

    raw = _normalize_line(raw)
    digits = re.sub(r'\D', '', raw)
    if len(digits) < 8:
        return None

    if len(digits) == 10:
        return digits

    if len(digits) == 11 and digits.startswith('0'):
        return digits[1:]

    if len(digits) == 11 and digits.startswith('1'):
        return digits[1:]

    if len(digits) == 12 and digits.startswith('91'):
        return digits[-10:]

    if len(digits) == 13 and digits.startswith('091'):
        return digits[-10:]

    if len(digits) > 10 and raw.startswith('+'):
        return digits[-10:]

    if has_label and len(digits) >= 8:
        return digits[-10:] if len(digits) > 10 else digits

    return None


def extract_mobile(text):
    """Extract phone/mobile number from resume text.
    Handles formats: 
    - 9876543210 (10 digits)
    - 91063 83952 (5-5 split)
    - +91 9876543210 / +91-9876543210
    - (123) 456-7890 / 123-456-7890 / 123.456.7890
    - +1 123 456 7890
    """
    normalized = _normalize_text(text)
    section_map = extract_section_map(normalized)
    contact_text = _section_text(section_map, 'contact')

    lines = []
    if contact_text:
        lines.extend([line for line in contact_text.split('\n') if line.strip()])
    lines.extend([line for line in normalized.split('\n')[:14] if line.strip()])
    lines.extend([line for line in normalized.split('\n') if re.search(r'phone|mobile|mob|cell|tel|contact|whatsapp', line, re.IGNORECASE)])
    lines.extend([line for line in normalized.split('\n') if len(re.sub(r'\D', '', line)) >= 5])

    phone_label_pattern = re.compile(r'\b(?:phone|mobile|mob|cell|tel|contact|whatsapp)\b', re.IGNORECASE)
    general_pattern = re.compile(r'(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)')
    grouped_patterns = [
        re.compile(r'\b\d{5}[\s.-]\d{5}\b'),
        re.compile(r'\b\d{4}[\s.-]\d{3}[\s.-]\d{3}\b'),
        re.compile(r'\b\d{3}[\s.-]\d{3}[\s.-]\d{4}\b'),
        re.compile(r'\b\d{2,5}[\s.-]\d{2,5}[\s.-]\d{2,5}\b'),
    ]

    candidate_lines = []
    for index, line in enumerate(lines):
        candidate_lines.append((line, phone_label_pattern.search(line) is not None))
        if index + 1 < len(lines):
            combined = f"{line} {lines[index + 1]}"
            if len(re.sub(r'\D', '', combined)) >= 8:
                candidate_lines.append((combined, phone_label_pattern.search(combined) is not None))

    for raw_line, has_label in candidate_lines:
        for pattern in grouped_patterns:
            for match in pattern.finditer(raw_line):
                candidate = _normalize_phone_candidate(match.group(0), has_label=has_label)
                if candidate:
                    return candidate

        for match in general_pattern.finditer(raw_line):
            candidate = _normalize_phone_candidate(match.group(0), has_label=has_label)
            if candidate:
                return candidate

    # Final fallback: scan the entire normalized text for any 10-digit run.
    for match in re.finditer(r'\b\d{10}\b', normalized):
        candidate = _normalize_phone_candidate(match.group(0))
        if candidate:
            return candidate

    return None


def parse_resume(path):
    text = ""
    if path.lower().endswith('.pdf'):
        text = extract_text_from_pdf(path)
    elif path.lower().endswith('.docx'):
        text = extract_text_from_docx(path)
    else:
        text = extract_text_from_pdf(path)

    skills = extract_skills(text)
    experience_duration = extract_experience_duration(text)
    name = extract_name(text)
    email = extract_email(text)
    mobile = extract_mobile(text)
    location = extract_location(text)

    # New structured extractions
    action_verbs = extract_action_verbs(text)
    metrics_data = extract_metrics_and_impact(text)
    experience_bullets = extract_experience_bullets(text)
    education = extract_education(text)
    formatting_score = check_formatting_quality(text)
    sections = extract_section_map(text)

    return {
        "text": text,
        "sections": sections,
        "skills": skills,
        "experience": experience_duration,
        "name": name,
        "email": email,
        "mobile": mobile,
        "location": location,
        "action_verbs": action_verbs,
        "metrics_data": metrics_data,
        "experience_bullets": experience_bullets,
        "education": education,
        "formatting_score": formatting_score,
        "extracted_keywords": [],
    }


def parse_job_description(text):
    text = text if isinstance(text, str) else text.decode('utf-8')
    skills = extract_skills(text)
    jd_keywords = extract_keywords_from_job(text)
    min_exp_matches = re.findall(r"(\d+)\s*(?:\+)?\s*(?:years|yrs|year|y)\b", text.lower())
    min_experience = max([int(v) for v in min_exp_matches], default=0)
    return {
        "text": text,
        "required_skills": skills,
        "jd_keywords": jd_keywords,
        "min_experience": min_experience
    }